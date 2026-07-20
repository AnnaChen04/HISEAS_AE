#!/usr/bin/env python3
"""Write need_to_confirm.docx and rejected_with_reason.docx."""
import json, os, math, collections, re
import openpyxl
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
flags = json.load(open(os.path.join(BASE, 'flags.json')))
rejects = json.load(open(os.path.join(BASE, 'rejects.json')))
corr = json.load(open(os.path.join(BASE, 'corrections_log.json')))

wb = openpyxl.load_workbook(os.path.join(BASE, 'SST_Data_Mining_corrected_2026-07-20.xlsx'), read_only=True)
o = openpyxl.load_workbook(os.path.join(BASE, 'SST_Data_Mining.xlsx'), read_only=True)
NEW = [s for s in wb.sheetnames if s not in set(o.sheetnames)]

def style(doc):
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = '' if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def bucket(flags, *pats):
    out = []
    for f in flags:
        if any(re.search(p, f['reason'], re.I) for p in pats):
            out.append(f)
    return out

# ===================================================================== FLAGS
doc = Document()
style(doc)
doc.add_heading('SST Data Mining — Items Needing Your Confirmation', 0)
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('NOAA NCEI Paleoclimatology advanced search, 140,000–115,000 cal yr BP, '
          'Global Ocean | Ocean, overlap-any. 122 studies returned; 485 data files examined. '
          'Generated 2026-07-20.')
doc.add_paragraph('Every item below was handled as described, but involves a judgement call '
                  'you should verify. Nothing here is silently discarded.')

# --- A. corrections to the user's own tabs
doc.add_heading('A. Corrections applied to YOUR existing tabs (copy only)', 1)
doc.add_paragraph('Applied to SST_Data_Mining_corrected_2026-07-20.xlsx. Your original '
                  'SST_Data_Mining.xlsx is untouched.')
table(doc, ['Tab', 'Field', 'Before', 'After', 'Basis'],
      [[c['tab'], c['field'], c['old'], c['new'], c['basis']] for c in corr],
      [0.9, 1.3, 1.4, 1.4, 2.1])
doc.add_paragraph('Outstanding: your V19-29 records latitude −3.25, whereas NCEI gives −3.57. '
                  'Not changed — please confirm which is correct.')

# --- B. interval-average products
b = bucket(flags, r'Interval-average')
doc.add_heading('B. Interval-average products — excluded from the spreadsheet', 1)
doc.add_paragraph('These report temperature averaged over an interval (e.g. "Early LIG maximum '
                  'annual temperature") rather than assigning an age to each SST value, so they '
                  'cannot be compared point-by-point against the reconstruction.')
table(doc, ['Core / site', 'Study', 'Data link'],
      [[f['core'], (f.get('study') or '')[:70], f['url']] for f in b], [1.0, 2.4, 3.3]) if b else \
      doc.add_paragraph('None encountered beyond those already rejected at screening.')

# --- C. sparse
b = bucket(flags, r'SPARSE')
doc.add_heading(f'C. Sparse records — fewer than 5 points in 115–140 ka ({len(b)})', 1)
doc.add_paragraph('Accepted into the spreadsheet, but low sample density may signal a poorly '
                  'constrained depth–age model.')
table(doc, ['Core / site', 'Detail', 'Data link'],
      [[f['core'], f['reason'], f['url']] for f in b], [1.1, 1.9, 3.7])

# --- D. uncertainty handling
b = bucket(flags, r'1sd DOUBLED', r'paired bounds', r'No uncertainty column')
doc.add_heading(f'D. Uncertainty (SST ±2sd) handling ({len(b)})', 1)
table(doc, ['Core / site', 'How ±2sd was derived', 'Data link'],
      [[f['core'], f['reason'], f['url']] for f in b], [1.1, 2.2, 3.4])

# --- E. proxy / calibration
b = bucket(flags, r'outside Hoffman', r'alternative', r'Seasonal')
doc.add_heading(f'E. Proxy type, calibration choice and seasonality ({len(b)})', 1)
doc.add_paragraph('Where a source offered several calibrations of one proxy, the first was used '
                  'and the alternatives are listed so you can override the choice.')
table(doc, ['Core / site', 'Issue', 'Data link'],
      [[f['core'], f['reason'][:190], f['url']] for f in b], [1.0, 2.9, 2.8])

# --- F. dedup edge cases
b = bucket(flags, r'DIFFERENT 1x1', r'not used', r'not written')
doc.add_heading(f'F. Duplicate / near-duplicate resolution ({len(b)})', 1)
doc.add_paragraph('The 1°×1° grid rule alone missed two cores whose coordinates straddle a cell '
                  'boundary; these were caught by core-ID matching and moved to the chronology file.')
table(doc, ['Core / site', 'Resolution', 'Data link'],
      [[f['core'], f['reason'][:190], f['url']] for f in b], [1.0, 2.9, 2.8])

# --- G. shared cells among accepted
def cell(lo, la):
    return (math.floor(lo) + .5, math.floor(la) + .5)
g = collections.defaultdict(list)
for t in NEW:
    g[cell(wb[t].cell(2, 2).value, wb[t].cell(2, 3).value)].append(t)
sh = [(k, v) for k, v in g.items() if len(v) > 1]
doc.add_heading('G. Newly accepted cores sharing a 1°×1° cell with each other', 1)
doc.add_paragraph('Not a rejection, but the RSOI combines same-cell cores by optimal estimation, '
                  'so these will be blended into one site.')
table(doc, ['Grid cell (lon, lat)', 'Tabs'], [[str(k), ', '.join(v)] for k, v in sh], [1.6, 5.0])

# --- H. structural
b = bucket(flags, r'HTML', r'Could not locate', r'Elevation', r'Non-point')
doc.add_heading(f'H. Structural / metadata problems ({len(b)})', 1)
table(doc, ['Core / site', 'Problem', 'Data link'],
      [[f['core'], f['reason'][:150], f['url']] for f in b], [1.1, 2.6, 3.0])

noproxy = [t for t in NEW if wb[t].cell(2, 4).value is None]
if noproxy:
    doc.add_paragraph('Proxy type could not be determined automatically for: '
                      + ', '.join(f'{t} ({wb[t].cell(2,23).value})' for t in noproxy))

doc.save(os.path.join(BASE, 'need_to_confirm.docx'))
print('need_to_confirm.docx  sections A-H, %d flagged items' % len(flags))

# ================================================================== REJECTS
doc = Document()
style(doc)
doc.add_heading('SST Data Mining — Rejected Studies and Reasons', 0)
doc.add_paragraph('Studies and data files that matched the search parameters but were not added '
                  'to the data-mining spreadsheet. Generated 2026-07-20.')

groups = collections.OrderedDict()
LABEL = [
    (r'no values inside', 'Contains SST but no data points within 115–140 ka'),
    (r'No SST column', 'No sea-surface-temperature column in the data table'),
    (r'No age column|no numeric', 'No age/time variable — depth only (see Jerry file where applicable)'),
    (r'Not marine', 'Not a marine ocean-core site'),
    (r'Download missing', 'Data file could not be downloaded'),
    (r'Could not locate|HTML', 'File is not a parseable data table'),
]
for f in rejects:
    lab = 'Other'
    for pat, name in LABEL:
        if re.search(pat, f['reason'], re.I):
            lab = name
            break
    groups.setdefault(lab, []).append(f)

doc.add_heading('Summary', 1)
table(doc, ['Reason', 'Files'], [[k, len(v)] for k, v in groups.items()] +
      [['TOTAL', sum(len(v) for v in groups.values())]], [4.5, 0.9])

for k, v in groups.items():
    doc.add_heading(f'{k} ({len(v)})', 1)
    table(doc, ['Core / site', 'Study', 'Data link'],
          [[f['core'], (f.get('study') or '')[:60], f['url']] for f in v], [1.0, 2.1, 3.6])

doc.save(os.path.join(BASE, 'rejected_with_reason.docx'))
print('rejected_with_reason.docx  %d rejected files in %d groups' % (len(rejects), len(groups)))
